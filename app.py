from flask import Flask, render_template, request, send_file
import pandas as pd
from bs4 import BeautifulSoup
import os
from docx import Document
from docx.shared import Pt
import tempfile
from docx2pdf import convert  # ✅ New import
import subprocess

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'

# Load student data
df = pd.read_csv("student.csv")
students = df.to_dict(orient="records")
student_map = dict(zip(df['ID'], df['NAME']))

@app.route("/")
def index():
    return render_template("index.html", students=students, student_map=student_map)

@app.route("/preview", methods=["POST"])
def preview():
    student_id = request.form["student_id"]
    student_name = request.form["student_name"]
    html_file = request.files["html_file"]

    if not html_file.filename.endswith(".html"):
        return "Only .html files are allowed"

    filename = os.path.join(app.config["UPLOAD_FOLDER"], html_file.filename)
    html_file.save(filename)

    with open(filename, "r", encoding="utf-8") as file:
        raw_html = file.read()

    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup.find_all(True):
        if tag.name.startswith("h"):
            tag['style'] = "font-family:Times New Roman; font-size:14pt;"
        else:
            tag['style'] = "font-family:Times New Roman; font-size:12pt;"

    formatted_html = str(soup)

    return render_template(
        "preview.html",
        student_id=student_id,
        student_name=student_name,
        html_code=raw_html,
        html_render=formatted_html
    )

@app.route("/download_word", methods=["POST"])
def download_word():
    student_id = request.form["student_id"]
    student_name = request.form["student_name"]
    html_code = request.form["html_code"]
    html_render = request.form["html_render"]

    doc = Document()
    doc.add_heading(f"{student_name} ({student_id})", level=0)

    doc.add_heading("HTML Code", level=1)
    code_paragraph = doc.add_paragraph()
    run = code_paragraph.add_run(html_code)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    doc.add_heading("Rendered Output", level=1)
    soup = BeautifulSoup(html_render, "html.parser")
    for tag in soup.find_all(['p', 'div', 'span', 'h1', 'h2', 'h3', 'h4', 'h5']):
        text = tag.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            p.style.font.name = "Times New Roman"
            p.style.font.size = Pt(12)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    return send_file(tmp_path, as_attachment=True, download_name="journal.docx")


@app.route("/download_pdf", methods=["POST"])
def download_pdf():
    student_id = request.form["student_id"]
    student_name = request.form["student_name"]
    html_code = request.form["html_code"]
    html_render = request.form["html_render"]

    # Generate Word file
    doc = Document()
    doc.add_heading(f"{student_name} ({student_id})", level=0)

    doc.add_heading("HTML Code", level=1)
    code_paragraph = doc.add_paragraph()
    run = code_paragraph.add_run(html_code)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    doc.add_heading("Rendered Output", level=1)
    soup = BeautifulSoup(html_render, "html.parser")
    for tag in soup.find_all(['p', 'div', 'span', 'h1', 'h2', 'h3', 'h4', 'h5']):
        text = tag.get_text(strip=True)
        if text:
            p = doc.add_paragraph(text)
            p.style.font.name = "Times New Roman"
            p.style.font.size = Pt(12)

    # Save DOCX to temp
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        docx_path = tmp_docx.name

    # Convert to PDF using LibreOffice
    pdf_path = docx_path.replace(".docx", ".pdf")
    subprocess.run([
        r"C:\Program Files\LibreOffice\program\soffice.exe",  # Update if path is different
        "--headless",
        "--convert-to", "pdf",
        "--outdir", os.path.dirname(docx_path),
        docx_path
    ], check=True)

    return send_file(pdf_path, as_attachment=True, download_name=f"{student_name}_journal.pdf")



if __name__ == "__main__":
    app.run(debug=True)
